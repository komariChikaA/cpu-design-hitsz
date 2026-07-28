#!/usr/bin/env python3
"""Insert readable datapath/state diagrams into the course DOCX template.

The source template keeps large portrait placeholders under each requirement.
This script replaces those placeholders with short navigation notes and inserts
the actual figures on dedicated A4 landscape pages.  Images remain inline so
Word and LibreOffice cannot float them over text.
"""

from __future__ import annotations

import argparse
import sys
from copy import deepcopy
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[2]
LOCAL_DEPS = REPO_ROOT / ".codex-deps" / "python"
if LOCAL_DEPS.exists():
    sys.path.insert(0, str(LOCAL_DEPS))

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_run_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), "Microsoft YaHei")
    r_fonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    r_fonts.set(qn("w:eastAsia"), "微软雅黑")


def set_placeholder_note(document, table_index: int, text: str) -> None:
    row = document.tables[table_index].rows[1]
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row.height = Inches(0.72)
    cell = row.cells[0]
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    set_run_font(paragraph.add_run(text), 10.5)


def find_body_paragraph(document, exact_text: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Cannot find report heading: {exact_text}")


def section_properties(base_sect_pr, landscape: bool):
    sect_pr = deepcopy(base_sect_pr)
    pg_sz = sect_pr.find(qn("w:pgSz"))
    if pg_sz is None:
        pg_sz = OxmlElement("w:pgSz")
        sect_pr.append(pg_sz)
    if landscape:
        pg_sz.set(qn("w:w"), "16839")
        pg_sz.set(qn("w:h"), "11907")
        pg_sz.set(qn("w:orient"), "landscape")
    else:
        pg_sz.set(qn("w:w"), "11907")
        pg_sz.set(qn("w:h"), "16839")
        pg_sz.attrib.pop(qn("w:orient"), None)

    pg_mar = sect_pr.find(qn("w:pgMar"))
    if pg_mar is None:
        pg_mar = OxmlElement("w:pgMar")
        sect_pr.append(pg_mar)
    for key, value in {
        "top": "737",
        "bottom": "737",
        "left": "737",
        "right": "737",
        "header": "340",
        "footer": "340",
        "gutter": "0",
    }.items():
        pg_mar.set(qn(f"w:{key}"), value)

    page_number = sect_pr.find(qn("w:pgNumType"))
    if page_number is not None:
        sect_pr.remove(page_number)

    section_type = sect_pr.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
        sect_pr.insert(0, section_type)
    section_type.set(qn("w:val"), "nextPage")
    return sect_pr


def section_break_element(sect_pr):
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    p_pr.append(sect_pr)
    paragraph.append(p_pr)
    return paragraph


def add_inline_figure(document, image_path: Path, caption: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(10.2))
    max_height = Inches(6.35)
    if shape.height > max_height:
        ratio = max_height / shape.height
        shape.height = max_height
        shape.width = int(shape.width * ratio)

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(2)
    caption_paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(caption_paragraph.add_run(caption), 10.5)
    return [paragraph._p, caption_paragraph._p]


def add_page_break(document):
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    return paragraph._p


def insert_landscape_figure_block(
    document,
    anchor_text: str,
    figures: list[tuple[Path, str]],
    base_sect_pr,
) -> None:
    anchor = find_body_paragraph(document, anchor_text)
    elements = [section_break_element(section_properties(base_sect_pr, False))]
    for index, (image_path, caption) in enumerate(figures):
        if index:
            elements.append(add_page_break(document))
        elements.extend(add_inline_figure(document, image_path, caption))
    elements.append(section_break_element(section_properties(base_sect_pr, True)))
    for element in elements:
        anchor._p.addprevious(element)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--template",
        type=Path,
        default=REPO_ROOT / "materials" / "计算机设计与实践-报告模板-20260723.docx",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=REPO_ROOT / "docs" / "course-report" / "计算机设计与实践-实验报告.docx",
    )
    args = parser.parse_args()

    evidence = REPO_ROOT / "docs" / "course-report" / "figures"
    detail = REPO_ROOT / "docs" / "datapath" / "report"
    required = [
        args.template,
        evidence / "01_singlecycle_datapath.png",
        evidence / "04_pipeline_datapath.png",
        evidence / "05_axi_state_machine.png",
        evidence / "06a_pipeline_load_use_hazard.png",
        evidence / "06b_no_cache_axi_read.png",
        evidence / "06c_no_cache_axi_write.png",
    ]
    missing = [path for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing report inputs: " + ", ".join(map(str, missing)))

    document = Document(str(args.template))
    base_sect_pr = deepcopy(document.sections[-1]._sectPr)

    set_placeholder_note(
        document,
        1,
        "完整的单周期数据通路总图见下一横向页；其后附模块级局部详图，便于核对接口信号名与位宽。",
    )
    set_placeholder_note(
        document,
        4,
        "完整的五级流水线数据通路总图见下一横向页；图中单独标出级间寄存器、前递与冒险控制。",
    )
    set_placeholder_note(
        document,
        5,
        "AXI 状态转换图及两组服务器仿真波形见后续横向页；波形分别展示流水线访存等待/恢复和无 Cache 的 AXI 读写事务。",
    )

    single_cycle_figures = [
        (
            evidence / "01_singlecycle_datapath.png",
            "图 1-1  miniRV 普通单周期 CPU 数据通路总图（无流水级寄存器）",
        ),
        (detail / "singlecycle-part-1-npc-pc.png", "图 1-2  IF：下一 PC 生成与 PC 寄存器"),
        (detail / "singlecycle-part-2-pc-rom.png", "图 1-3  IF：PC 与指令存储器"),
        (detail / "singlecycle-part-3-decode.png", "图 1-4  ID：指令字段拆分"),
        (detail / "singlecycle-part-4-control-sext.png", "图 1-5  ID：控制器与立即数扩展"),
        (detail / "singlecycle-part-5-writeback-rf.png", "图 1-6  WB/ID：写回选择与寄存器堆"),
        (detail / "singlecycle-part-6-ex.png", "图 1-7  EX：操作数选择与 ALU"),
        (detail / "singlecycle-part-7-mreq-dram.png", "图 1-8  MEM：访存请求与数据存储器"),
        (detail / "singlecycle-part-8-dram-mext.png", "图 1-9  MEM/WB：数据存储器与读数据扩展"),
    ]
    insert_landscape_figure_block(
        document,
        "1.2 单周期CPU详细设计",
        single_cycle_figures,
        base_sect_pr,
    )
    insert_landscape_figure_block(
        document,
        "2.2 流水线SoC详细设计",
        [
            (
                evidence / "04_pipeline_datapath.png",
                "图 2-1  miniRV 五级流水线 CPU 数据通路总图",
            )
        ],
        base_sect_pr,
    )
    insert_landscape_figure_block(
        document,
        "2.3 流水线SoC仿真分析",
        [
            (
                evidence / "05_axi_state_machine.png",
                "图 2-2  无 Cache 的 AXI4 单事务主机状态转换图",
            ),
            (
                evidence / "06a_pipeline_load_use_hazard.png",
                "图 2-3  流水线 Load-Use 测试中的访存等待、气泡与恢复波形",
            ),
            (
                evidence / "06b_no_cache_axi_read.png",
                "图 2-4  无 Cache 的 AXI 读事务及 AR/R 通道握手波形",
            ),
            (
                evidence / "06c_no_cache_axi_write.png",
                "图 2-5  无 Cache 的 AXI 写事务及 AW/W/B 通道握手波形",
            ),
        ],
        base_sect_pr,
    )

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(args.output))
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
